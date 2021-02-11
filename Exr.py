import os  # импорт необходимых библиотек для работы
import docx
import re


try:
  from PIL.ExifTags import TAGS
except ImportError:
  TAGS = {}
try:
    from exceptions import PendingDeprecationWarning
except ImportError:
    pass

path = '/Users/user/Desktop/Extensions 2020'  # расположение файлов

text_data = [
    {'old_text': 'Folta', 'new_text': 'Danielewicz'},
    {'old_text': 'Dusan', 'new_text': 'Olgierd'},
    {'old_text': 'US $6,600 + VAT and an invoice will be sent', 'new_text': 'waived'},
    {'old_text': 'Brand President', 'new_text': 'Chief Operations Officer'},
    {'old_text': '30.07.2020', 'new_text': '25.01.2021'},
    {'old_text': 'US $6,600', 'new_text': '0'}]  # определяет переменные для замены

old_date = re.compile('dd.\s([a-zA-Z0-9\s]*)\sreferred to')  # позволяет вытащить дату документа
new_date = '25 January 2021'  # новая дата документа
today = '11 February 2021'  # новая дата для SCHEDULE B

with os.scandir(path) as listOfEntries:  # вытаскивает из папки размещения список файлов
    for entry in listOfEntries:  # начинает цикл для списка файлов
        if entry.is_file():
            file_name = path + '\\' + entry.name  # определяет имя файла для работы с указанием пути к файлу
            document = docx.Document(file_name)  # определяет формат файла
            document.paragraphs[0].text = new_date  # заменяет в первом параграфе (абзаце) дату на новую
            for paragraph in document.paragraphs:  # для каждого параграфа (абзаца) в документе
                if re.search(old_date, paragraph.text):  # ищет паттерн в тексте и возвращает true
                    search_result = re.search(old_date, paragraph.text).group(1)  # ищет дату по тексту
                    paragraph.text = paragraph.text.replace(search_result, today)  # подменяет дату в тексте
                for element in text_data:
                    paragraph.text = paragraph.text.replace(element['old_text'], element[
                        'new_text'])  # тут понятно просто меняет старый текст на новый
            for table in document.tables:  # здесь та же проверка, но не для абзацев, а для таблиц
                for row in table.rows:  # берет ряды
                    for cell in row.cells:  # проверяет ячейки в рядах
                        for element in text_data:  # для текстового элемента
                            cell.text = cell.text.replace(element['old_text'], element['new_text'])  # заменяет текст
            document.save(file_name)  # сохраняет документ